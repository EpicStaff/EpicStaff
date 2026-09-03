import csv
import zlib

from pdfminer.ascii85 import ascii85decode, asciihexdecode

import pdfplumber
from docx import Document
from io import BytesIO, StringIO
from bs4 import BeautifulSoup
from loguru import logger

from utils.extraction_limits import (
    ExtractionBudget,
    ExtractionLimitExceeded,
    default_budget,
)


def extract_text_from_binary(
    binary_content: bytes, file_type: str, budget: ExtractionBudget | None = None
) -> str:
    """Universal dispatcher to extract text from binary content based on file type."""

    file_type = file_type.lower().lstrip(".")
    budget = budget or default_budget()
    budget.check_input_size(binary_content)

    try:
        if file_type in ("txt", "md", "json"):
            return extract_text(binary_content)

        elif file_type == "pdf":
            return extract_text_from_pdf(binary_content, budget=budget)

        elif file_type == "csv":
            return extract_text_from_csv(binary_content)

        elif file_type == "docx":
            return extract_text_from_docx(binary_content, budget=budget)

        elif file_type == "html":
            return extract_text_from_html(binary_content, budget=budget)

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    except Exception as e:
        logger.error("Failed to extract text from {} file: {}", file_type, e)
        raise


def extract_text(binary_content: bytes) -> str:
    """Extract text from plain text files."""

    try:
        return binary_content.decode("utf-8")
    except UnicodeDecodeError:
        logger.warning("UTF-8 decode failed, trying latin-1")
        return binary_content.decode("latin-1")


def _is_valid_pdf(binary_content: bytes) -> bool:
    """
    Check if binary content is a valid PDF by looking for the PDF magic bytes.
    PDF files should start with '%PDF-' (possibly with leading whitespace).
    """
    # Strip leading whitespace and check for PDF signature
    content = binary_content.lstrip()
    return content.startswith(b"%PDF-")


def _extract_page_text(page) -> str:
    """
    Make pdfplumber extraction output similar to PyMuPDF output.
    """
    lines = page.extract_text_lines(
        strip=True,
        return_chars=False,
        x_tolerance=3,
        y_tolerance=3,
    )
    if not lines:
        return ""
    if len(lines) == 1:
        return lines[0]["text"]

    spacings = [
        lines[i]["top"] - lines[i - 1]["top"]
        for i in range(1, len(lines))
        if lines[i]["top"] > lines[i - 1]["top"]
    ]
    if not spacings:
        return "\n".join(line["text"] + " " for line in lines).rstrip(" ")

    median_spacing = sorted(spacings)[len(spacings) // 2]
    paragraph_threshold = median_spacing * 1.5

    parts = [lines[0]["text"] + " "]
    for i in range(1, len(lines)):
        spacing = lines[i]["top"] - lines[i - 1]["top"]
        if spacing > paragraph_threshold:
            parts.append(" ")
        parts.append(lines[i]["text"] + " ")

    return "\n".join(parts).rstrip(" ")


# ASCII85 and ASCIIHex only ever shrink their input, so neither can hide a bomb;
# deflate is the one filter that amplifies, and it is inflated through the cap.
# Anything else (LZW, RunLength) is rare and cannot be bounded, so it is refused.
_SHRINKING_FILTERS = {
    "ASCII85Decode": ascii85decode,
    "A85": ascii85decode,
    "ASCIIHexDecode": asciihexdecode,
    "AHx": asciihexdecode,
}
_FLATE_FILTERS = ("FlateDecode", "Fl")


def _filter_names(resolved) -> list[str]:
    """Return a stream's /Filter entries in the order the parser will apply them."""
    declared = resolved.attrs.get("Filter")
    return [
        getattr(f, "name", str(f))
        for f in (declared if isinstance(declared, list) else [declared])
        if f is not None
    ]


def _bounded_stream_size(raw: bytes, names: list[str], limit: int) -> int:
    """Return what a stream decodes to, never buffering more than limit + 1 bytes."""
    data = raw

    # The chain must be walked in order: the real ratio sits behind the last
    # filter, so measuring an intermediate layer would wave a bomb through.
    for name in names:
        if name in _FLATE_FILTERS:
            try:
                data = zlib.decompressobj().decompress(data, limit + 1)
            except zlib.error:
                # Corrupt stream: let pdfplumber report it, do not mask it here
                return len(data)
            if len(data) > limit:
                return len(data)
        elif name in _SHRINKING_FILTERS:
            try:
                data = _SHRINKING_FILTERS[name](data)
            except Exception:
                return len(data)
        else:
            raise ExtractionLimitExceeded(
                f"Page content uses the filter {name!r}, whose decoded size "
                "cannot be bounded"
            )

    return len(data)


def _decoded_content_size(page, limit: int) -> int:
    """Return the decoded size of a page's content streams, bounded by limit."""
    # A PDF declares no decoded size -- /Length is the compressed count -- so
    # the only way to learn it is to decode, bounded, before layout happens.
    contents = getattr(page.page_obj, "contents", None)
    streams = contents if isinstance(contents, list) else [contents]
    total = 0

    for stream in streams:
        resolved = stream.resolve() if hasattr(stream, "resolve") else stream
        raw = getattr(resolved, "rawdata", None) or b""
        total += _bounded_stream_size(raw, _filter_names(resolved), limit - total)
        if total > limit:
            break

    return total


def extract_text_from_pdf(
    binary_content: bytes, budget: ExtractionBudget | None = None
) -> str:
    """Extract text from PDF files, falling back to plain text for invalid PDFs."""

    budget = budget or default_budget()

    if not _is_valid_pdf(binary_content):
        logger.warning(
            "Content has .pdf extension but is not a valid PDF file. "
            "Attempting plain text extraction."
        )
        return extract_text(binary_content)

    text_parts = []
    try:
        with pdfplumber.open(BytesIO(binary_content)) as pdf_document:
            for page in pdf_document.pages:
                budget.count_page()
                budget.add_content_bytes(
                    _decoded_content_size(page, budget.max_content_bytes)
                )
                page_text = _extract_page_text(page)
                if page_text:
                    text_parts.append(page_text)

        if not text_parts:
            logger.warning("No text extracted from PDF")
            return ""

        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error("PDF text extraction failed: {}", e)
        raise


def extract_text_from_csv(binary_content: bytes) -> str:
    """Extract text from CSV files."""

    try:
        text_content = binary_content.decode("utf-8")
        csv_file = StringIO(text_content)

        delimiter = ","
        reader = csv.reader(csv_file, delimiter=delimiter)

        extracted_rows = []
        for row in reader:
            if row and len(row[0].replace(delimiter, "")) != 0:
                extracted_rows.append(",".join(row))

        return "\n".join(extracted_rows)

    except Exception as e:
        logger.error("CSV text extraction failed: {}", e)
        raise


def extract_text_from_docx(
    binary_content: bytes, budget: ExtractionBudget | None = None
) -> str:
    """Extract text from DOCX files using python-docx."""

    budget = budget or default_budget()
    # Document() materialises the whole package, so this must run before it
    budget.check_unpacked_size(binary_content)

    try:
        docx_file = BytesIO(binary_content)
        document = Document(docx_file)

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        if not paragraphs:
            logger.warning("No text extracted from DOCX")
            return ""

        return "\n".join(paragraphs)

    except Exception as e:
        logger.error("DOCX text extraction failed: {}", e)
        raise


def extract_text_from_html(
    binary_content: bytes, budget: ExtractionBudget | None = None
) -> str:
    """Extract text from HTML files using BeautifulSoup, dropping scripts, styles and images."""

    budget = budget or default_budget()
    budget.check_html_size(binary_content)

    try:
        html_content = binary_content.decode("utf-8")
        soup = BeautifulSoup(html_content, "html.parser")

        # Remove unwanted tags
        for tag in soup(["script", "style", "img"]):
            tag.decompose()

        # Return cleaned HTML
        return str(soup)

    except Exception as e:
        logger.error("HTML text extraction failed: {}", e)
        raise
