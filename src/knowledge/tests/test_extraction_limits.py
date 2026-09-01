"""Text extraction is bounded in input size, unpacked size and page count.

The extractor runs on the knowledge worker's ThreadPoolExecutor behind
`indexing_semaphore = 3` (main.py), so three files that never finish extracting
stall ingestion for every tenant. A PDF can declare thousands of pages, and
pdfplumber walks all of them synchronously with no way to interrupt.

Each cap is asserted through the public extract_* entry points rather than
against ExtractionBudget alone -- a budget nothing consults would satisfy a
unit test of the budget itself while leaving the worker just as pinnable.
"""

import time
import zlib
import zipfile
from io import BytesIO

import pytest
from docx import Document

from utils.extraction_limits import (
    ExtractionBudget,
    ExtractionLimitExceeded,
    default_budget,
)
from utils.file_text_extractor import (
    extract_text_from_binary,
    extract_text_from_csv,
    extract_text_from_docx,
    extract_text_from_pdf,
)


def build_pdf(page_texts: list[str]) -> bytes:
    """Build a minimal single-font PDF carrying one text line per entry."""
    objects: dict[int, bytes] = {}
    n_pages = len(page_texts)
    page_ids = [3 + 2 * i for i in range(n_pages)]
    content_ids = [4 + 2 * i for i in range(n_pages)]
    font_id = 3 + 2 * n_pages

    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    kids = b" ".join(b"%d 0 R" % page_id for page_id in page_ids)
    objects[2] = b"<< /Type /Pages /Kids [%s] /Count %d >>" % (kids, n_pages)

    for index, text in enumerate(page_texts):
        objects[page_ids[index]] = (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] "
            b"/Contents %d 0 R /Resources << /Font << /F1 %d 0 R >> >> >>"
            % (content_ids[index], font_id)
        )
        stream = b"BT /F1 12 Tf 20 200 Td (%s) Tj ET" % text.encode("ascii")
        objects[content_ids[index]] = b"<< /Length %d >>\nstream\n%s\nendstream" % (
            len(stream),
            stream,
        )

    objects[font_id] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

    body = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for number in sorted(objects):
        offsets[number] = len(body)
        body += b"%d 0 obj\n%s\nendobj\n" % (number, objects[number])

    xref_offset = len(body)
    last = max(objects)
    body += b"xref\n0 %d\n0000000000 65535 f \n" % (last + 1)
    for number in range(1, last + 1):
        body += b"%010d 00000 n \n" % offsets.get(number, 0)
    body += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (
        last + 1,
        xref_offset,
    )
    return bytes(body)


def build_pdf_bomb(decoded_bytes: int) -> bytes:
    """Build a one-page PDF whose FlateDecode content stream decodes hugely."""
    chunk = b"BT /F1 12 Tf 20 200 Td (" + b"A" * 500 + b") Tj ET\n"
    plain = chunk * max(1, decoded_bytes // len(chunk))
    compressed = zlib.compress(plain, 9)

    objects = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        4: b"<< /Length %d /Filter /FlateDecode >>\nstream\n" % len(compressed)
        + compressed
        + b"\nendstream",
        5: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }

    body = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for number in sorted(objects):
        offsets[number] = len(body)
        body += b"%d 0 obj\n%s\nendobj\n" % (number, objects[number])

    xref_offset = len(body)
    last = max(objects)
    body += b"xref\n0 %d\n0000000000 65535 f \n" % (last + 1)
    for number in range(1, last + 1):
        body += b"%010d 00000 n \n" % offsets.get(number, 0)
    body += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (
        last + 1,
        xref_offset,
    )
    return bytes(body)


def build_docx(paragraphs: list[str]) -> bytes:
    """Build an in-memory DOCX carrying one paragraph per entry."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_CONTENT_TYPES = (
    b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    b'<Default Extension="xml" ContentType="application/xml"/>'
    b'<Default Extension="rels" ContentType="application/vnd.openxml'
    b'formats-package.relationships+xml"/>'
    b'<Override PartName="/word/document.xml" ContentType="application/vnd.openxml'
    b'formats-officedocument.wordprocessingml.document.main+xml"/></Types>'
)

_RELS = (
    b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/'
    b'relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org'
    b'/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
    b"</Relationships>"
)


def build_docx_bomb(unpacked_bytes: int) -> bytes:
    """Build a small DOCX whose document.xml declares `unpacked_bytes` unpacked."""
    paragraph = b"<w:p><w:r><w:t>" + b"A" * 900 + b"</w:t></w:r></w:p>"
    head = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:document xmlns:w="http://schemas.openxmlformats.org/'
        b'wordprocessingml/2006/main"><w:body>'
    )
    tail = b"</w:body></w:document>"
    count = max(1, (unpacked_bytes - len(head) - len(tail)) // len(paragraph))

    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        archive.writestr("[Content_Types].xml", _CONTENT_TYPES)
        archive.writestr("_rels/.rels", _RELS)
        archive.writestr("word/document.xml", head + paragraph * count + tail)
    return buffer.getvalue()


def permissive_budget(**overrides) -> ExtractionBudget:
    """Build a budget whose caps are all effectively unlimited unless overridden."""
    limits = {
        "max_input_bytes": 10_000_000,
        "max_unpacked_bytes": 10_000_000,
        "max_content_bytes": 10_000_000,
        "max_html_bytes": 10_000_000,
        "max_pages": 10_000,
    }
    limits.update(overrides)
    return ExtractionBudget(**limits)


# --- page cap ---


def test_pdf_extraction_rejects_more_pages_than_cap():
    pdf_bytes = build_pdf([f"Page number {i}" for i in range(5)])

    with pytest.raises(ExtractionLimitExceeded, match="page"):
        extract_text_from_pdf(pdf_bytes, budget=permissive_budget(max_pages=2))


def test_pdf_extraction_accepts_page_count_at_the_cap():
    pdf_bytes = build_pdf([f"Page number {i}" for i in range(5)])

    text = extract_text_from_pdf(pdf_bytes, budget=permissive_budget(max_pages=5))

    assert "Page number 0" in text
    assert "Page number 4" in text


def test_pdf_page_cap_trips_before_every_page_is_parsed():
    """The cap must bound work done, not just reject after walking the whole file."""
    pdf_bytes = build_pdf([f"Page number {i}" for i in range(200)])
    budget = permissive_budget(max_pages=3)

    with pytest.raises(ExtractionLimitExceeded):
        extract_text_from_pdf(pdf_bytes, budget=budget)

    assert budget.pages_seen == 4


# --- decoded page content cap ---
#
# PDF declares no decompressed size: /Length is the *compressed* byte count, so
# unlike a DOCX there is no header to read. The cost is also not in the inflate
# -- it is pdfplumber laying out every glyph, ~24s per decoded MB. So each
# page's content stream is inflated through a hard cap and measured before any
# layout runs.


def test_pdf_extraction_rejects_a_page_decoding_over_the_content_cap():
    bomb = build_pdf_bomb(1_000_000)
    assert len(bomb) < 10_000  # tiny on the wire

    with pytest.raises(ExtractionLimitExceeded, match="content"):
        extract_text_from_pdf(bomb, budget=permissive_budget(max_content_bytes=100_000))


def test_pdf_content_cap_rejects_before_layout_runs():
    """Layout is the expensive part; rejecting after it would defeat the cap."""
    bomb = build_pdf_bomb(1_000_000)  # ~24s if it reaches layout

    started = time.perf_counter()
    with pytest.raises(ExtractionLimitExceeded):
        extract_text_from_pdf(bomb, budget=permissive_budget(max_content_bytes=100_000))
    elapsed = time.perf_counter() - started

    assert elapsed < 5.0, f"took {elapsed:.1f}s -- layout ran before the cap fired"


def test_pdf_content_cap_is_enforced_through_the_dispatcher():
    bomb = build_pdf_bomb(1_000_000)

    with pytest.raises(ExtractionLimitExceeded, match="content"):
        extract_text_from_binary(
            bomb, "pdf", budget=permissive_budget(max_content_bytes=100_000)
        )


def test_content_cap_accepts_an_ordinary_pdf():
    pdf_bytes = build_pdf(["Ordinary content", "Second page"])

    text = extract_text_from_pdf(pdf_bytes, budget=permissive_budget())

    assert "Ordinary content" in text


def test_content_budget_accumulates_across_pages():
    """No single page is oversized; only their sum is."""
    pdf_bytes = build_pdf([f"Page number {i}" for i in range(20)])
    budget = permissive_budget(max_content_bytes=200)

    with pytest.raises(ExtractionLimitExceeded, match="content"):
        extract_text_from_pdf(pdf_bytes, budget=budget)


# --- unpacked size cap ---
#
# A DOCX is a ZIP. python-docx parses the whole package inside Document(),
# before any paragraph is reachable, so an output-side cap is too late: the
# 5 MB fixture below peaks at ~10 MB of resident XML before the first
# paragraph exists. The declared entry sizes are readable from the ZIP
# central directory without decompressing anything, so the check runs first.


def test_docx_extraction_rejects_a_document_that_unpacks_over_the_cap():
    bomb = build_docx_bomb(5_000_000)
    assert len(bomb) < 100_000  # small on the wire, huge unpacked

    with pytest.raises(ExtractionLimitExceeded, match="unpacks to"):
        extract_text_from_docx(
            bomb, budget=permissive_budget(max_unpacked_bytes=100_000)
        )


def test_docx_unpacked_cap_is_enforced_through_the_dispatcher():
    bomb = build_docx_bomb(5_000_000)

    with pytest.raises(ExtractionLimitExceeded, match="unpacks to"):
        extract_text_from_binary(
            bomb, "docx", budget=permissive_budget(max_unpacked_bytes=100_000)
        )


def test_unpacked_cap_accepts_an_ordinary_docx():
    docx_bytes = build_docx(["first paragraph", "second paragraph"])

    text = extract_text_from_docx(docx_bytes, budget=permissive_budget())

    assert text == "first paragraph\nsecond paragraph"


def test_unpacked_check_ignores_input_that_is_not_a_zip():
    """PDF, CSV and plain text are not container formats; the check must pass them."""
    budget = permissive_budget(max_unpacked_bytes=10)

    budget.check_unpacked_size(b"just some plain bytes, not a zip at all")


# --- extraction still works ---


def test_docx_extraction_returns_every_paragraph():
    """The DOCX loop was rewritten when the char cap went; it still reads all text."""
    docx_bytes = build_docx(["first paragraph", "second paragraph"])

    assert extract_text_from_docx(docx_bytes) == "first paragraph\nsecond paragraph"


def test_csv_extraction_returns_every_row():
    csv_bytes = b"a,b,c\nd,e,f"

    assert extract_text_from_csv(csv_bytes) == "a,b,c\nd,e,f"


# --- input byte cap ---


def test_extraction_rejects_input_over_byte_cap_before_parsing():
    html_bytes = b"<html><body>" + b"x" * 5000 + b"</body></html>"

    with pytest.raises(ExtractionLimitExceeded, match="bytes"):
        extract_text_from_binary(
            html_bytes, "html", budget=permissive_budget(max_input_bytes=1000)
        )


def test_input_byte_cap_applies_to_every_file_type():
    payload = b"z" * 5000

    with pytest.raises(ExtractionLimitExceeded, match="bytes"):
        extract_text_from_binary(
            payload, "txt", budget=permissive_budget(max_input_bytes=1000)
        )


# --- integration with the dispatcher and defaults ---


def test_dispatcher_threads_budget_through_to_pdf_extraction():
    pdf_bytes = build_pdf([f"Page number {i}" for i in range(5)])

    with pytest.raises(ExtractionLimitExceeded, match="page"):
        extract_text_from_binary(
            pdf_bytes, "pdf", budget=permissive_budget(max_pages=2)
        )


def test_default_budget_extracts_an_ordinary_document_unchanged():
    pdf_bytes = build_pdf(["Ordinary content", "Second page"])

    text = extract_text_from_binary(pdf_bytes, "pdf")

    assert "Ordinary content" in text
    assert "Second page" in text


def test_default_budget_caps_are_finite():
    budget = default_budget()

    assert 0 < budget.max_pages < 1_000_000
    assert 0 < budget.max_input_bytes < 1_000_000_000
    assert 0 < budget.max_unpacked_bytes < 1_000_000_000
    assert 0 < budget.max_content_bytes < 1_000_000_000


def test_caps_match_the_upload_path():
    """django_app must refuse at upload anything this service would refuse.

    Separate containers, so the numbers cannot be imported across -- they are
    asserted on both sides. See
    django_app/tests/services_tests/test_document_upload_limits.py
    ::test_caps_match_the_knowledge_service.
    """
    budget = default_budget()

    assert budget.max_input_bytes == 50 * 1024 * 1024
    assert budget.max_unpacked_bytes == 256 * 1024 * 1024


def test_limit_error_is_a_value_error_so_existing_handlers_still_catch_it():
    """execute_preview_chunking reports ValueError as a `failed` job, not a crash."""
    assert issubclass(ExtractionLimitExceeded, ValueError)
